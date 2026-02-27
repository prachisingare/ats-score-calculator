"""
File parser module for extracting text from PDF and DOCX files.

Goals (production-focused):
- Stable, clean text output for keyword matching + embeddings
- Better multi-column PDF extraction (pdfplumber first, PyPDF2 fallback)
- DOCX extraction without duplicate/scrambled text
- Preserve resume structure: section breaks + bullet lines
- Detect image-based PDFs (no extractable text)
"""

import logging
import re
from typing import BinaryIO, List, Optional

from PyPDF2 import PdfReader
from docx import Document

# Try to import pdfplumber for better multi-column support
try:
    import pdfplumber  # type: ignore
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False

logger = logging.getLogger(__name__)


class FileParseError(Exception):
    """Exception raised when file parsing fails."""
    pass


# Common bullet glyphs found in PDFs/DOCX
_BULLETS = {"•", "●", "◦", "▪", "–", "—", "·", "‣", "∙", "○", "□", "■"}
_NUM_BULLET_RE = re.compile(r"^\s*(\d+[\)\.]|[a-zA-Z][\)\.])\s+")


def _clean_extracted_text(text: str) -> str:
    """
    Clean up extracted text while preserving structure (important for resumes).

    - Fix hyphenated line breaks: "machine-\\nlearning" -> "machine-learning"
    - Normalize bullets to "- " (glyph bullets + numbered bullets)
    - Collapse excessive spaces/tabs but keep newlines
    - Keep at most one blank line between blocks
    """
    if not text:
        return ""

    # Normalize newlines
    text = text.replace("\r\n", "\n").replace("\r", "\n")

    # Fix hyphenated line breaks
    text = re.sub(r"(\w)-\n(\w)", r"\1-\2", text)

    # Collapse horizontal whitespace (keep newlines)
    text = re.sub(r"[ \t]+", " ", text)

    lines = text.split("\n")
    cleaned_lines: List[str] = []
    prev_blank = False

    for raw in lines:
        line = raw.strip()

        # If bullet appears inside the line, split it to preserve list structure
        # Example: "Skills • Python • SQL" -> "Skills" + "- Python" + "- SQL"
        if "•" in line and not line.startswith("•"):
            parts = [p.strip() for p in line.split("•") if p.strip()]
            if len(parts) >= 2:
                cleaned_lines.append(parts[0])
                for p in parts[1:]:
                    cleaned_lines.append("- " + p)
                prev_blank = False
                continue

        # Normalize glyph bullets
        if line and line[0] in _BULLETS:
            line = "- " + line[1:].strip()

        # Normalize numbered bullets: "1) foo" / "1. foo" / "a) foo"
        if line and _NUM_BULLET_RE.match(line):
            line = "- " + _NUM_BULLET_RE.sub("", line).strip()

        # Keep at most one blank line (preserve structure)
        if not line:
            if cleaned_lines and not prev_blank:
                cleaned_lines.append("")
            prev_blank = True
            continue

        cleaned_lines.append(line)
        prev_blank = False

    # Remove trailing blank lines
    while cleaned_lines and cleaned_lines[-1] == "":
        cleaned_lines.pop()

    return "\n".join(cleaned_lines).strip()


# ---------------------------
# PDF helpers
# ---------------------------
def _words_to_lines(words: list[dict], y_tol: float = 3.0) -> str:
    """
    Build lines from pdfplumber extract_words() output.
    This is a stronger fallback when extract_text() scrambles multi-columns.
    """
    if not words:
        return ""

    # Sort by top (y), then x
    words_sorted = sorted(words, key=lambda w: (round(w.get("top", 0) / y_tol), w.get("x0", 0)))

    lines: List[List[str]] = []
    current_key = None
    current: List[str] = []

    for w in words_sorted:
        top_bucket = round(w.get("top", 0) / y_tol)
        key = top_bucket
        txt = (w.get("text") or "").strip()
        if not txt:
            continue

        if current_key is None:
            current_key = key

        if key != current_key:
            if current:
                lines.append(current)
            current = [txt]
            current_key = key
        else:
            current.append(txt)

    if current:
        lines.append(current)

    # Join words into lines
    joined_lines = [" ".join(line).strip() for line in lines if line]
    return "\n".join(joined_lines).strip()


def _extract_text_with_pdfplumber(file: BinaryIO) -> str:
    """
    Extract text using pdfplumber (better for multi-column layouts).

    Strategy:
    1) page.extract_text(use_text_flow=True)
    2) if weak/empty, fallback to page.extract_words() and rebuild lines
    """
    text_parts: List[str] = []

    logger.info(f"Extracting PDF with pdfplumber (available: {PDFPLUMBER_AVAILABLE})")

    with pdfplumber.open(file) as pdf:
        logger.debug(f"PDF has {len(pdf.pages)} pages")
        for i, page in enumerate(pdf.pages):
            page_text = page.extract_text(
                x_tolerance=3,
                y_tolerance=3,
                use_text_flow=True,
            ) or ""

            cleaned = _clean_extracted_text(page_text)

            # If extract_text() is weak, rebuild from words (multi-column safer)
            if len(cleaned) < 80:
                try:
                    words = page.extract_words(
                        x_tolerance=2,
                        y_tolerance=2,
                        keep_blank_chars=False,
                        use_text_flow=True,
                    ) or []
                    rebuilt = _words_to_lines(words, y_tol=3.0)
                    rebuilt_clean = _clean_extracted_text(rebuilt)
                    if len(rebuilt_clean) > len(cleaned):
                        cleaned = rebuilt_clean
                except Exception as e:
                    logger.debug(f"Page {i+1}: extract_words fallback failed: {e}")

            if cleaned:
                text_parts.append(cleaned)
                logger.debug(f"Page {i+1}: extracted {len(cleaned)} chars")
            else:
                logger.debug(f"Page {i+1}: no text extracted")

    result = "\n\n".join(text_parts).strip()
    logger.info(f"pdfplumber extracted {len(result)} chars total")
    return result


def _extract_text_with_pypdf2(file: BinaryIO) -> str:
    """Extract text using PyPDF2 (fallback)."""
    logger.info("Extracting PDF with PyPDF2 (fallback)")
    reader = PdfReader(file)
    text_parts: List[str] = []

    for i, page in enumerate(reader.pages):
        page_text = page.extract_text() or ""
        cleaned = _clean_extracted_text(page_text)
        if cleaned:
            text_parts.append(cleaned)
            logger.debug(f"Page {i+1}: extracted {len(cleaned)} chars")

    result = "\n\n".join(text_parts).strip()
    logger.info(f"PyPDF2 extracted {len(result)} chars total")
    return result


def extract_text_from_pdf(file: BinaryIO) -> str:
    """
    Extract text content from a PDF file.

    Tries pdfplumber first (better for multi-column layouts),
    falls back to PyPDF2 if unavailable or returns empty text.
    """
    logger.info("Starting PDF text extraction")

    try:
        text = ""

        if PDFPLUMBER_AVAILABLE:
            try:
                file.seek(0)
                text = _extract_text_with_pdfplumber(file)
                if text.strip():
                    logger.info(f"PDF extraction successful: {len(text)} chars")
                    return text
                else:
                    logger.warning("pdfplumber returned empty text, trying PyPDF2")
            except Exception as e:
                logger.warning(f"pdfplumber failed: {e}, falling back to PyPDF2")

        file.seek(0)
        text = _extract_text_with_pypdf2(file)

        if not text.strip():
            # very likely image-based PDF
            raise FileParseError("PDF appears image-based or contains no extractable text")

        logger.info(f"PDF extraction successful (PyPDF2): {len(text)} chars")
        return text

    except FileParseError:
        raise
    except Exception as e:
        logger.error(f"PDF extraction failed: {e}")
        raise FileParseError(f"Failed to parse PDF: {str(e)}")


# ---------------------------
# DOCX helpers
# ---------------------------
def _extract_text_from_xml_element(element) -> str:
    """Fallback XML extraction when normal DOCX extraction is too short."""
    parts: List[str] = []

    if getattr(element, "text", None) and element.text and element.text.strip():
        parts.append(element.text.strip())

    for child in element:
        child_text = _extract_text_from_xml_element(child)
        if child_text:
            parts.append(child_text)
        if getattr(child, "tail", None) and child.tail and child.tail.strip():
            parts.append(child.tail.strip())

    return " ".join(parts)


def extract_text_from_docx(file: BinaryIO) -> str:
    """
    Extract text content from a DOCX file.

    Production-safe behavior:
    - Primary: paragraphs + tables + header/footer (keeps order stable)
    - De-duplicates repeated table cells per row
    - Fallback: XML extraction ONLY if primary result is too short
    """
    logger.info("Starting DOCX text extraction")

    try:
        document = Document(file)
        parts: List[str] = []

        # Paragraphs
        para_count = 0
        for paragraph in document.paragraphs:
            t = (paragraph.text or "").strip()
            if t:
                parts.append(t)
                para_count += 1
        logger.debug(f"Extracted {para_count} paragraphs")

        # Tables (dedupe cell text)
        table_rows = 0
        for table in document.tables:
            for row in table.rows:
                row_cells: List[str] = []
                seen = set()
                for cell in row.cells:
                    ct = (cell.text or "").strip()
                    if ct and ct not in seen:
                        row_cells.append(ct)
                        seen.add(ct)
                if row_cells:
                    parts.append(" | ".join(row_cells))
                    table_rows += 1
        logger.debug(f"Extracted {table_rows} table rows")

        # Header/Footer
        header_footer_count = 0
        for section in document.sections:
            for block in (section.header, section.footer):
                try:
                    for p in block.paragraphs:
                        t = (p.text or "").strip()
                        if t:
                            parts.append(t)
                            header_footer_count += 1
                except Exception:
                    pass
        logger.debug(f"Extracted {header_footer_count} header/footer items")

        text = "\n".join(parts).strip()
        logger.info(f"Primary DOCX extraction: {len(text)} chars")

        # Fallback to XML if too short (often indicates text boxes / complex layout)
        if len(text) < 200:
            logger.warning(f"Primary extraction too short ({len(text)} chars), trying XML fallback")
            try:
                body = document.element.body
                xml_text = _extract_text_from_xml_element(body).strip()
                if len(xml_text) > len(text):
                    logger.info(f"XML fallback extracted {len(xml_text)} chars")
                    text = xml_text
            except Exception as e:
                logger.warning(f"XML fallback failed: {e}")

        if not text.strip():
            raise FileParseError("DOCX appears empty or contains no extractable text")

        cleaned = _clean_extracted_text(text)
        logger.info(f"DOCX extraction successful: {len(cleaned)} chars after cleaning")
        return cleaned

    except FileParseError:
        raise
    except Exception as e:
        raise FileParseError(f"Failed to parse DOCX: {str(e)}")


def _is_docx_file(file: BinaryIO) -> bool:
    """Check if file is actually a DOCX (Office Open XML) format."""
    try:
        file.seek(0)
        header = file.read(4)
        file.seek(0)
        return header == b"PK\x03\x04"
    except Exception:
        return False


def extract_text_from_file(file: BinaryIO, filename: str) -> str:
    """
    Extract text from a file based on its extension.

    Supported formats: .pdf, .docx, .txt
    """
    logger.info(f"Extracting text from file: {filename}")
    filename_lower = filename.lower()

    if filename_lower.endswith(".pdf"):
        return extract_text_from_pdf(file)

    if filename_lower.endswith(".docx"):
        if not _is_docx_file(file):
            logger.error(f"File {filename} appears to be renamed .doc file")
            raise FileParseError(
                "This appears to be a legacy .doc file renamed as .docx. "
                "Please convert to proper .docx format or .pdf"
            )
        return extract_text_from_docx(file)

    if filename_lower.endswith(".doc"):
        raise FileParseError("Legacy .doc format is not supported. Please convert to .docx or .pdf")

    if filename_lower.endswith(".txt"):
        try:
            content = file.read()
            if isinstance(content, bytes):
                try:
                    return content.decode("utf-8")
                except UnicodeDecodeError:
                    return content.decode("latin-1", errors="ignore")
            return str(content)
        except Exception as e:
            raise FileParseError(f"Failed to read text file: {str(e)}")

    raise FileParseError("Unsupported file type. Supported formats: .pdf, .docx, .txt")


def get_supported_extensions() -> list[str]:
    """Return list of supported file extensions."""
    return [".pdf", ".docx", ".txt"]